from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(filename="Smart_Tution_Fee_Portal_Documentation.docx"):
    doc = Document()

    # Title
    title = doc.add_heading('Smart Tution Fee Portal Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Abstract
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "The Smart Tution Fee Portal is a web-based platform designed to simplify "
        "and automate the collection, management, and tracking of tuition fees "
        "for educational institutions or private coaching centers. It provides a "
        "centralized database for administrators to easily monitor payments, pending dues, "
        "and generate financial reports. For students and parents, the portal offers a "
        "convenient and secure way to pay fees online, view payment history, and receive "
        "timely notifications for upcoming dues. The system reduces manual paperwork, "
        "minimizes errors, and enhances transparency in financial transactions."
    )

    # 2. System Comparison
    doc.add_heading('2. Comparison of New System with Existing System', level=1)
    doc.add_heading('Existing System', level=2)
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run('Manual record keeping in registers or basic excel sheets.')
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run('Prone to human errors during calculation.')
    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run('Time-consuming receipt generation.')
    p4 = doc.add_paragraph(style='List Bullet')
    p4.add_run('Difficult to track defaulters and pending dues.')

    doc.add_heading('New System', level=2)
    p5 = doc.add_paragraph(style='List Bullet')
    p5.add_run('Automated digital records and secure cloud database.')
    p6 = doc.add_paragraph(style='List Bullet')
    p6.add_run('Error-free automated calculations.')
    p7 = doc.add_paragraph(style='List Bullet')
    p7.add_run('Instant digital receipts and online secure payments.')
    p8 = doc.add_paragraph(style='List Bullet')
    p8.add_run('Automated reminders and instant defaulter list generation.')

    # 3. Requirements
    doc.add_heading('3. Hardware and Software Requirement Specification', level=1)
    doc.add_heading('Software Requirements:', level=2)
    p9 = doc.add_paragraph(style='List Bullet')
    p9.add_run('Frontend: HTML5, CSS3, JavaScript (React/Bootstrap)')
    p10 = doc.add_paragraph(style='List Bullet')
    p10.add_run('Backend: Python (Django/Flask) or Node.js / PHP')
    p11 = doc.add_paragraph(style='List Bullet')
    p11.add_run('Database: MySQL / PostgreSQL / MongoDB')
    p12 = doc.add_paragraph(style='List Bullet')
    p12.add_run('Server: Apache / Nginx / Node server')

    doc.add_heading('Hardware Requirements:', level=2)
    p13 = doc.add_paragraph(style='List Bullet')
    p13.add_run('Processor: Intel Core i3 or above')
    p14 = doc.add_paragraph(style='List Bullet')
    p14.add_run('RAM: 4 GB or higher')
    p15 = doc.add_paragraph(style='List Bullet')
    p15.add_run('Storage: 50 GB free space')
    p16 = doc.add_paragraph(style='List Bullet')
    p16.add_run('Internet Connection for online payments')

    # 4. Modules
    doc.add_heading('4. Modules and their short description', level=1)
    p17 = doc.add_paragraph(style='List Bullet')
    p17.add_run('Admin Module: ').bold = True
    p17.add_run('Manage students, teachers, courses, fee structures, and view financial reports.')
    p18 = doc.add_paragraph(style='List Bullet')
    p18.add_run('Student/Parent Module: ').bold = True
    p18.add_run('View fee details, pay online, download receipts, check payment history.')
    p19 = doc.add_paragraph(style='List Bullet')
    p19.add_run('Payment Gateway Integration: ').bold = True
    p19.add_run('Secure processing of online transactions.')
    p20 = doc.add_paragraph(style='List Bullet')
    p20.add_run('Notification Module: ').bold = True
    p20.add_run('Send automated SMS/Email reminders for pending fees.')
    p21 = doc.add_paragraph(style='List Bullet')
    p21.add_run('Reporting Module: ').bold = True
    p21.add_run('Generate daily/monthly revenue and defaulters reports.')

    # 5. Users
    doc.add_heading('5. Users and their role description', level=1)
    p22 = doc.add_paragraph(style='List Bullet')
    p22.add_run('Admin: ').bold = True
    p22.add_run('Has full control over the system. Can add/remove students, configure fees, and view all reports.')
    p23 = doc.add_paragraph(style='List Bullet')
    p23.add_run('Student/Parent: ').bold = True
    p23.add_run('Can log in to view their specific fee dues, make payments, and download receipts.')
    p24 = doc.add_paragraph(style='List Bullet')
    p24.add_run('Accountant (Optional): ').bold = True
    p24.add_run('Can manage day-to-day transactions and approve manual payments if any.')

    # 6. Diagrams
    doc.add_heading('6. Diagrams', level=1)
    doc.add_paragraph('Please insert the generated diagrams (from the provided markdown or images) into the sections below:')
    doc.add_heading('6.1 Data Flow Diagrams (DFD)', level=2)
    doc.add_paragraph('[Insert Context Level DFD (Level 0) here]')
    doc.add_paragraph('[Insert Level 1 DFDs here]')
    doc.add_paragraph('[Insert Level 2 DFDs here]')
    
    doc.add_heading('6.2 Use Case Diagram', level=2)
    doc.add_paragraph('[Insert Use Case Diagrams here]')
    
    doc.add_heading('6.3 Class Diagram', level=2)
    doc.add_paragraph('[Insert Class Diagram here]')
    
    doc.add_heading('6.4 Activity Diagram', level=2)
    doc.add_paragraph('[Insert Activity Diagrams here]')
    
    doc.add_heading('6.5 E-R Diagram', level=2)
    doc.add_paragraph('[Insert E-R Diagram here]')
    
    # 7. E-R Description
    doc.add_heading('7. Description of E-R Diagram', level=1)
    doc.add_paragraph(
        "This ER (Entity Relationship) Diagram represents the model of Smart Tution Fee Portal. "
        "The entity-relationship diagram shows all the visual instruments of database tables and the relations "
        "between Admin, Teacher, Students, Courses, Fee Transactions, and Notifications. It uses structured data "
        "to define the relationships between these entities.\n\n"
        "The main entities of the system are Admin, Student, Course, Fee_Structure, Payment_Transaction, and Notification. "
        "The system revolves around Admin and Students entities. The Admin can manage courses, fee structures, "
        "and student records. Students belong to specific courses, which determines their fee structure. "
        "Students can perform Payment Transactions, generating digital receipts. The Notification entity handles "
        "alerts for pending and paid dues linked to students."
    )

    doc.save(filename)
    print(f"Documentation saved as {filename}")

if __name__ == '__main__':
    create_docx()
