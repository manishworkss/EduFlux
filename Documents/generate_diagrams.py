import os
import subprocess
from docx import Document
from docx.shared import Inches

os.makedirs("/Users/manishkumar/Desktop/tution/mermaid_tmp", exist_ok=True)

# Define all diagrams
diagrams_data = {
    "1_context.mmd": ("Figure 5.1.1. Context Level DFD: 0 Level", """graph TD
    A[Admin] -->|Manage Students, Fees| B((EduFlux System))
    B -->|Analytics & Reports| A
    C[Student] -->|Pay Fees, View History| B
    B -->|Receipts, Reminders| C
    B <-->|Process Payments| D[Razorpay Gateway]
    B <-->|AI Insights| E[Gemini AI]"""),

    "2_dfd_1_admin.mmd": ("Figure 5.2.1. Admin Side DFD: 1st Level", """graph TD
    Admin -->|Login Credentials| 1.0(Authentication)
    1.0 -->|Token| Admin
    Admin -->|Student Details| 2.0(Manage Students)
    2.0 -->|Save Data| DB[(Database)]
    Admin -->|Fee Configurations| 3.0(Manage Fee Structures)
    3.0 -->|Save Config| DB
    Admin -->|Assign Fees| 4.0(Assign to Students)
    4.0 -->|Update Due Dates| DB
    DB -->|Read Data| 5.0(Generate Analytics)
    5.0 -->|View Charts| Admin"""),

    "3_dfd_1_student.mmd": ("Figure 5.2.2. Student Side DFD: 1st Level", """graph TD
    Student -->|Login Credentials| 1.0(Authentication)
    1.0 -->|Token| Student
    Student -->|Request Dashboard| 2.0(Fetch Pending Fees)
    DB[(Database)] -->|Return Dues| 2.0
    2.0 -->|Display Dues| Student
    Student -->|Initiate Payment| 3.0(Process Payment)
    3.0 <-->|Gateway Comm| Razorpay
    3.0 -->|Save Transaction| DB
    DB -->|Fetch History| 4.0(Generate Receipt)
    4.0 -->|Download PDF| Student"""),

    "4_dfd_2_fees.mmd": ("Figure 5.3.1. Admin Side DFD: Manage Fees 2nd Level", """graph TD
    Admin -->|Input Fee Info| 3.1(Validate Inputs)
    3.1 -->|Valid Info| 3.2(Create Fee Record)
    3.2 -->|Store Record| DB[(Database)]
    DB -->|Fetch Templates| 3.3(List Fee Structures)
    3.3 -->|Display| Admin"""),

    "5_dfd_2_payment.mmd": ("Figure 5.3.2. Student Side DFD: Payment Process 2nd Level", """graph TD
    Student -->|Select Fee| 3.1(Create Order)
    3.1 -->|Order ID| Razorpay
    Razorpay -->|Order Created| 3.2(Render Checkout)
    3.2 -->|Enter Card/UPI| Razorpay
    Razorpay -->|Payment Signature| 3.3(Verify Payment)
    3.3 -->|Update Status| DB[(Database)]
    3.3 -->|Success/Failure| Student"""),

    "6_dfd_2_login.mmd": ("Figure 5.3.3. DFD: Login Process 2nd Level", """graph TD
    User -->|Email, Password| 1.1(Validate Credentials)
    1.1 -->|Query User| DB[(Database)]
    DB -->|Hash Match| 1.2(Generate JWT Token)
    1.2 -->|Return Token| User"""),

    "7_uc_admin.mmd": ("Figure 6.0.1. Admin Use Case Diagram", """flowchart LR
    Admin([Admin])
    subgraph EduFlux System
        UC1(Login / Logout)
        UC2(Manage Students)
        UC3(Create Fee Structures)
        UC4(Assign Fees)
        UC5(View Payments)
        UC6(View Analytics Dashboard)
        UC7(Chat with AI Assistant)
    end
    Admin --> UC1
    Admin --> UC2
    Admin --> UC3
    Admin --> UC4
    Admin --> UC5
    Admin --> UC6
    Admin --> UC7"""),

    "8_uc_student.mmd": ("Figure 6.0.2. Student Use Case Diagram", """flowchart LR
    Student([Student])
    subgraph EduFlux System
        UC1(Login / Logout)
        UC2(View Assigned Fees)
        UC3(Pay Fees Online)
        UC4(Download Receipts)
        UC5(View Payment History)
        UC6(Chat with AI Assistant)
        UC7(Razorpay)
    end
    Student --> UC1
    Student --> UC2
    Student --> UC3
    Student --> UC4
    Student --> UC5
    Student --> UC6
    UC3 -.-> UC7"""),

    "9_class_diagram.mmd": ("Figure 7.0.1. Class Diagram", """classDiagram
    class User {
        +Long id
        +String name
        +String email
        +String password
        +String role
        +login()
    }
    class Student {
        +Long id
        +String phone
        +Long adminId
    }
    class FeeStructure {
        +Long id
        +String name
        +Double amount
        +String frequency
    }
    class StudentFee {
        +Long id
        +Date dueDate
        +String status
    }
    class Payment {
        +Long id
        +Double amountPaid
        +Date paymentDate
        +String razorpayPaymentId
        +String razorpayOrderId
    }
    User <|-- Student
    Student "1" -- "*" StudentFee : has
    FeeStructure "1" -- "*" StudentFee : mapped to
    StudentFee "1" -- "*" Payment : triggers"""),

    "10_act_admin.mmd": ("Figure 8.0.1. Admin Activity Diagram", """stateDiagram-v2
    [*] --> Login
    Login --> Dashboard : Valid
    Dashboard --> SelectMenu
    SelectMenu --> ManageStudents
    SelectMenu --> SetupFee
    SelectMenu --> Analytics
    SetupFee --> FillFeeDetails
    FillFeeDetails --> SaveFee : Submit
    SaveFee --> [*]"""),

    "11_act_student.mmd": ("Figure 8.0.2. Student Activity Diagram", """stateDiagram-v2
    [*] --> Login
    Login --> StudentDashboard : Valid
    StudentDashboard --> CheckPendingDues
    CheckPendingDues --> ClickPayNow
    ClickPayNow --> Gateway : Redirect
    Gateway --> Success : Verified
    Gateway --> Failed : Failed
    Success --> DownloadReceipt
    DownloadReceipt --> [*]
    Failed --> CheckPendingDues"""),

    "12_er_diagram.mmd": ("Figure 9.0.1. E-R Diagram", """erDiagram
    ADMIN ||--o{ STUDENT : manages
    ADMIN ||--o{ FEE_STRUCTURE : creates
    FEE_STRUCTURE ||--o{ STUDENT_FEE : assigned_as
    STUDENT ||--o{ STUDENT_FEE : owes
    STUDENT_FEE ||--o{ PAYMENT : has
    STUDENT {
        bigint id PK
        string email
        string name
        string phone
        string role
        bigint admin_id FK
    }
    FEE_STRUCTURE {
        bigint id PK
        string name
        double amount
        string frequency
        bigint admin_id FK
    }
    STUDENT_FEE {
        bigint id PK
        bigint student_id FK
        bigint fee_structure_id FK
        date due_date
        string status
    }
    PAYMENT {
        bigint id PK
        bigint student_fee_id FK
        double amount_paid
        date payment_date
        string razorpay_id
        string razorpay_order_id
        string status
    }""")
}

os.chdir("/Users/manishkumar/Desktop/tution/mermaid_tmp")

# Write MMD files and generate PNGs
for filename, (title, content) in diagrams_data.items():
    with open(filename, "w") as f:
        f.write(content)
    
    png_file = filename.replace(".mmd", ".png")
    print(f"Generating {png_file}...")
    subprocess.run(["npx", "-p", "@mermaid-js/mermaid-cli", "mmdc", "-i", filename, "-o", png_file], check=True)

print("All PNGs generated. Creating Docx...")

doc = Document()
doc.add_heading('EduFlux System Diagrams', 0)

for filename, (title, content) in diagrams_data.items():
    png_file = filename.replace(".mmd", ".png")
    
    # Add title with keep_with_next so it stays with the picture
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    runner = p.add_run(title)
    runner.bold = True
    
    try:
        # Use a much smaller width (4.0 inches) so it fits easily on the page
        p_img = doc.add_paragraph()
        p_img.add_run().add_picture(f'/Users/manishkumar/Desktop/tution/mermaid_tmp/{png_file}', width=Inches(4.0))
        # Add some space after the picture
        doc.add_paragraph("")
    except Exception as e:
        doc.add_paragraph(f"Error loading image: {str(e)}")

doc.save('/Users/manishkumar/Desktop/tution/EduFlux_Diagrams_Full.docx')
print("Successfully generated EduFlux_Diagrams_Full.docx")
