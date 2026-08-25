import os
import subprocess
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Mermaid Definitions ---

mermaid_definitions = {
    "dfd_context": """
graph TD
    A[Admin] -->|Manage Fees & Students| B((Smart Tuition Fee Portal))
    C[Student/Parent] -->|Pay Fees & Check Dues| B
    B -->|Send Receipts & Reminders| C
    B -->|Generate Reports| A
    """,

    "dfd_level1_admin": """
graph TD
    A[Admin] -->|Login| B(1.0 Login System)
    B -->|Success| C(2.0 Dashboard)
    C -->|Add/Update| D(3.0 Manage Students)
    C -->|Set Fees| E(4.0 Manage Fee Structure)
    C -->|View| F(5.0 Generate Reports)
    D --> DB1[(Database)]
    E --> DB1
    F -.->|Fetch Data| DB1
    """,

    "dfd_level1_student": """
graph TD
    A[Student] -->|Login| B(1.0 Login System)
    B -->|Success| C(2.0 Student Dashboard)
    C -->|View Dues| D(3.0 Fee Details)
    C -->|Pay| E(4.0 Payment Gateway)
    E -->|Success| F(5.0 Download Receipt)
    D -.-> DB1[(Database)]
    E -->|Record Transaction| DB1
    """,

    "use_case": """
usecaseDiagram
    actor Admin
    actor Student
    Admin --> (Login)
    Admin --> (Manage Students)
    Admin --> (Manage Fee Structure)
    Admin --> (View Reports)
    Student --> (Login)
    Student --> (View Fee Dues)
    Student --> (Make Payment)
    Student --> (Download Receipt)
    (Make Payment) ..> (Verify Transaction) : <<include>>
    """,

    "class_diagram": """
classDiagram
    class Admin {
        +int admin_id
        +String name
        +String email
        +login()
        +manageStudents()
        +generateReports()
    }
    class Student {
        +int student_id
        +String name
        +String batch
        +float total_fee
        +login()
        +viewDues()
        +payFee()
    }
    class Payment {
        +int payment_id
        +float amount
        +Date date
        +String status
        +processPayment()
        +generateReceipt()
    }
    class Course {
        +int course_id
        +String course_name
        +float fee
    }
    Admin "1" -- "*" Student : manages
    Student "*" -- "1" Course : enrolled in
    Student "1" -- "*" Payment : makes
    """,

    "activity_diagram_student": """
stateDiagram-v2
    [*] --> Login
    Login --> VerifyCredentials
    VerifyCredentials --> Invalid: Fail
    Invalid --> Login
    VerifyCredentials --> Dashboard: Success
    Dashboard --> ViewDues
    ViewDues --> MakePayment
    MakePayment --> ProcessingGateway
    ProcessingGateway --> PaymentSuccess
    ProcessingGateway --> PaymentFailed
    PaymentFailed --> MakePayment
    PaymentSuccess --> GenerateReceipt
    GenerateReceipt --> [*]
    """,

    "er_diagram": """
erDiagram
    ADMIN ||--o{ STUDENT : manages
    ADMIN ||--|{ COURSE : configures
    STUDENT }|--|| COURSE : "enrolled in"
    STUDENT ||--o{ PAYMENT : makes
    PAYMENT ||--|| RECEIPT : generates
    
    STUDENT {
        int student_id PK
        string name
        string email
        string phone
    }
    COURSE {
        int course_id PK
        string name
        float total_fees
    }
    PAYMENT {
        int payment_id PK
        float amount
        date payment_date
        string status
    }
    ADMIN {
        int admin_id PK
        string username
        string password
    }
    """
}

def generate_images():
    for name, definition in mermaid_definitions.items():
        mmd_file = f"{name}.mmd"
        png_file = f"{name}.png"
        
        # fix use case syntax which mermaid doesn't support directly, let's use a standard graph for use case
        if name == "use_case":
            definition = """
graph LR
    Admin([Admin])
    Student([Student])
    L(Login)
    MS(Manage Students)
    MF(Manage Fees)
    VR(View Reports)
    VD(View Dues)
    MP(Make Payment)
    DR(Download Receipt)
    Admin --- L
    Admin --- MS
    Admin --- MF
    Admin --- VR
    Student --- L
    Student --- VD
    Student --- MP
    Student --- DR
            """

        with open(mmd_file, "w") as f:
            f.write(definition.strip())
            
        print(f"Generating image for {name}...")
        subprocess.run(["npx", "@mermaid-js/mermaid-cli", "-i", mmd_file, "-o", png_file, "-b", "white"], check=True)

def create_docx(filename="Smart_Tution_Fee_Portal_Documentation.docx"):
    doc = Document()

    title = doc.add_heading('Smart Tution Fee Portal Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "The Smart Tution Fee Portal is a web-based platform designed to simplify "
        "and automate the collection, management, and tracking of tuition fees "
        "for educational institutions or private coaching centers. It provides a "
        "centralized database for administrators to easily monitor payments, pending dues, "
        "and generate financial reports. For students and parents, the portal offers a "
        "convenient and secure way to pay fees online, view payment history, and receive "
        "timely notifications for upcoming dues. The system reduces manual paperwork, "
        "minimizes errors, and enhances transparency in financial transactions."
    )

    doc.add_heading('2. Comparison of New System with Existing System', level=1)
    doc.add_heading('Existing System', level=2)
    doc.add_paragraph('Manual record keeping in registers or basic excel sheets.', style='List Bullet')
    doc.add_paragraph('Prone to human errors during calculation.', style='List Bullet')
    doc.add_paragraph('Time-consuming receipt generation.', style='List Bullet')
    doc.add_paragraph('Difficult to track defaulters and pending dues.', style='List Bullet')

    doc.add_heading('New System', level=2)
    doc.add_paragraph('Automated digital records and secure cloud database.', style='List Bullet')
    doc.add_paragraph('Error-free automated calculations.', style='List Bullet')
    doc.add_paragraph('Instant digital receipts and online secure payments.', style='List Bullet')
    doc.add_paragraph('Automated reminders and instant defaulter list generation.', style='List Bullet')

    doc.add_heading('3. Hardware and Software Requirement Specification', level=1)
    doc.add_heading('Software Requirements:', level=2)
    doc.add_paragraph('Frontend: HTML5, CSS3, JavaScript (React/Bootstrap)', style='List Bullet')
    doc.add_paragraph('Backend: Python (Django/Flask) or Node.js / PHP', style='List Bullet')
    doc.add_paragraph('Database: MySQL / PostgreSQL / MongoDB', style='List Bullet')
    doc.add_paragraph('Server: Apache / Nginx / Node server', style='List Bullet')
    doc.add_heading('Hardware Requirements:', level=2)
    doc.add_paragraph('Processor: Intel Core i3 or above', style='List Bullet')
    doc.add_paragraph('RAM: 4 GB or higher', style='List Bullet')
    doc.add_paragraph('Storage: 50 GB free space', style='List Bullet')
    doc.add_paragraph('Internet Connection for online payments', style='List Bullet')

    doc.add_heading('4. Modules and their short description', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Admin Module: ').bold = True
    p.add_run('Manage students, teachers, courses, fee structures, and view financial reports.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Student/Parent Module: ').bold = True
    p.add_run('View fee details, pay online, download receipts, check payment history.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Payment Gateway Integration: ').bold = True
    p.add_run('Secure processing of online transactions.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Notification Module: ').bold = True
    p.add_run('Send automated SMS/Email reminders for pending fees.')

    doc.add_heading('5. Users and their role description', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Admin: ').bold = True
    p.add_run('Has full control over the system. Can add/remove students, configure fees, and view reports.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Student/Parent: ').bold = True
    p.add_run('Can log in to view their specific fee dues, make payments, and download receipts.')

    doc.add_heading('6. Diagrams', level=1)
    
    doc.add_heading('6.1 Data Flow Diagrams (DFD)', level=2)
    doc.add_paragraph('Context Level DFD (Level 0):')
    if os.path.exists("dfd_context.png"):
        doc.add_picture("dfd_context.png", width=Inches(5))
        
    doc.add_paragraph('Level 1 DFD (Admin):')
    if os.path.exists("dfd_level1_admin.png"):
        doc.add_picture("dfd_level1_admin.png", width=Inches(5))

    doc.add_paragraph('Level 1 DFD (Student):')
    if os.path.exists("dfd_level1_student.png"):
        doc.add_picture("dfd_level1_student.png", width=Inches(5))

    doc.add_heading('6.2 Use Case Diagram', level=2)
    if os.path.exists("use_case.png"):
        doc.add_picture("use_case.png", width=Inches(5))
        
    doc.add_heading('6.3 Class Diagram', level=2)
    if os.path.exists("class_diagram.png"):
        doc.add_picture("class_diagram.png", width=Inches(5))
        
    doc.add_heading('6.4 Activity Diagram', level=2)
    if os.path.exists("activity_diagram_student.png"):
        doc.add_picture("activity_diagram_student.png", width=Inches(4))
        
    doc.add_heading('6.5 E-R Diagram', level=2)
    if os.path.exists("er_diagram.png"):
        doc.add_picture("er_diagram.png", width=Inches(5))

    doc.add_heading('7. Description of E-R Diagram', level=1)
    doc.add_paragraph(
        "This ER (Entity Relationship) Diagram represents the model of Smart Tution Fee Portal. "
        "The entity-relationship diagram shows all the visual instruments of database tables and the relations "
        "between Admin, Teacher, Students, Courses, Fee Transactions, and Notifications. It uses structured data "
        "to define the relationships between these entities.\n\n"
        "The main entities of the system are Admin, Student, Course, Fee_Structure, Payment_Transaction, and Notification. "
        "The system revolves around Admin and Students entities. The Admin can manage courses, fee structures, "
        "and student records. Students belong to specific courses, which determines their fee structure. "
        "Students can perform Payment Transactions, generating digital receipts."
    )

    doc.save(filename)
    print(f"Documentation saved as {filename}")

if __name__ == '__main__':
    try:
        generate_images()
    except Exception as e:
        print("Error generating images:", e)
    create_docx()
